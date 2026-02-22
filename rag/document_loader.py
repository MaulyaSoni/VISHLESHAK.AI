"""
Document Loader for Vishleshak AI v1
Loads and processes documents from various formats into vector store
"""

from typing import List, Dict, Any, Optional
from pathlib import Path
import logging
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from config import rag_config

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Process documents for vector store ingestion
    
    Supports:
    - Text files (.txt)
    - Markdown (.md)
    - PDF (.pdf)
    - Word documents (.docx)
    - CSV (.csv)
    - JSON (.json)
    
    Features:
    - Smart text chunking
    - Metadata extraction
    - Multiple file format support
    
    Usage:
        processor = DocumentProcessor()
        documents = processor.load_directory("knowledge_base/education")
        chunks = processor.chunk_documents(documents)
    """
    
    def __init__(self):
        """Initialize document processor"""
        self.chunk_size = rag_config.CHUNK_SIZE
        self.chunk_overlap = rag_config.CHUNK_OVERLAP
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        logger.info("✅ Document processor initialized")
    
    def load_text_file(self, file_path: Path) -> Document:
        """
        Load a text file
        
        Args:
            file_path: Path to text file
            
        Returns:
            Document object with content and metadata
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "text"
            }
            
            return Document(page_content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Error loading text file {file_path}: {e}")
            return None
    
    def load_markdown_file(self, file_path: Path) -> Document:
        """Load a Markdown file"""
        # Markdown is just text, so use text loader
        doc = self.load_text_file(file_path)
        if doc:
            doc.metadata["type"] = "markdown"
        return doc
    
    def load_pdf_file(self, file_path: Path) -> List[Document]:
        """
        Load a PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            List of Document objects (one per page)
        """
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            documents = []
            
            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    metadata = {
                        "source": str(file_path),
                        "filename": file_path.name,
                        "type": "pdf",
                        "page": page_num + 1,
                        "total_pages": len(reader.pages)
                    }
                    
                    documents.append(Document(page_content=text, metadata=metadata))
            
            logger.info(f"Loaded {len(documents)} pages from PDF: {file_path.name}")
            return documents
        
        except Exception as e:
            logger.error(f"Error loading PDF file {file_path}: {e}")
            return []
    
    def load_docx_file(self, file_path: Path) -> Document:
        """
        Load a Word document
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Document object
        """
        try:
            from docx import Document as DocxDocument
            
            doc = DocxDocument(file_path)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            content = "\n\n".join(paragraphs)
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "docx",
                "paragraphs": len(paragraphs)
            }
            
            return Document(page_content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Error loading DOCX file {file_path}: {e}")
            return None
    
    def load_csv_file(self, file_path: Path) -> Document:
        """
        Load a CSV file as document
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            Document object with CSV content
        """
        try:
            import pandas as pd
            
            df = pd.read_csv(file_path)
            
            # Convert to text representation
            content = f"CSV File: {file_path.name}\n\n"
            content += f"Columns: {', '.join(df.columns)}\n\n"
            content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
            content += "Sample Data:\n"
            content += df.head(10).to_string()
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "csv",
                "rows": len(df),
                "columns": len(df.columns)
            }
            
            return Document(page_content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Error loading CSV file {file_path}: {e}")
            return None
    
    def load_json_file(self, file_path: Path) -> Document:
        """
        Load a JSON file
        
        Args:
            file_path: Path to JSON file
            
        Returns:
            Document object
        """
        try:
            import json
            
            with open(file_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
            
            # Convert to readable text
            content = json.dumps(data, indent=2)
            
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "type": "json"
            }
            
            return Document(page_content=content, metadata=metadata)
        
        except Exception as e:
            logger.error(f"Error loading JSON file {file_path}: {e}")
            return None
    
    def load_file(self, file_path: Path) -> List[Document]:
        """
        Load a file of any supported type
        
        Args:
            file_path: Path to file
            
        Returns:
            List of Document objects
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        logger.info(f"Loading file: {file_path.name}")
        
        # Route to appropriate loader
        if extension == ".txt":
            doc = self.load_text_file(file_path)
            return [doc] if doc else []
        
        elif extension == ".md":
            doc = self.load_markdown_file(file_path)
            return [doc] if doc else []
        
        elif extension == ".pdf":
            return self.load_pdf_file(file_path)
        
        elif extension == ".docx":
            doc = self.load_docx_file(file_path)
            return [doc] if doc else []
        
        elif extension == ".csv":
            doc = self.load_csv_file(file_path)
            return [doc] if doc else []
        
        elif extension == ".json":
            doc = self.load_json_file(file_path)
            return [doc] if doc else []
        
        else:
            logger.warning(f"Unsupported file format: {extension}")
            return []
    
    def load_directory(
        self, 
        directory_path: Path,
        recursive: bool = True
    ) -> List[Document]:
        """
        Load all supported documents from a directory
        
        Args:
            directory_path: Path to directory
            recursive: Whether to search subdirectories
            
        Returns:
            List of all loaded documents
        """
        directory_path = Path(directory_path)
        
        if not directory_path.exists():
            logger.error(f"Directory does not exist: {directory_path}")
            return []
        
        logger.info(f"Loading documents from: {directory_path}")
        
        documents = []
        pattern = "**/*" if recursive else "*"
        
        for file_path in directory_path.glob(pattern):
            if file_path.is_file():
                if file_path.suffix.lower() in rag_config.KNOWLEDGE_FILE_EXTENSIONS:
                    docs = self.load_file(file_path)
                    documents.extend(docs)
        
        logger.info(f"Loaded {len(documents)} documents from directory")
        return documents
    
    def chunk_documents(self, documents: List[Document]) -> List[Document]:
        """
        Split documents into chunks for vector storage
        
        Args:
            documents: List of documents to chunk
            
        Returns:
            List of chunked documents
        """
        if not documents:
            return []
        
        logger.info(f"Chunking {len(documents)} documents...")
        
        chunked_docs = self.text_splitter.split_documents(documents)
        
        logger.info(f"Created {len(chunked_docs)} chunks")
        return chunked_docs
    
    def process_documents_for_vectorstore(
        self, 
        documents: List[Document]
    ) -> tuple[List[str], List[Dict]]:
        """
        Process documents into format for vector store
        
        Args:
            documents: List of documents
            
        Returns:
            Tuple of (texts, metadatas)
        """
        texts = [doc.page_content for doc in documents]
        metadatas = [doc.metadata for doc in documents]
        
        return texts, metadatas
